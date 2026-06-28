"""
📦 BAZA PUBLISHER
=================
Har yangi test yaratilganda Baza Guruhiga avtomatik e'lon qilish.

Jarayon:
  1. Test yaratildi (bot/sayt/fayl/quiz forward — qayerdan bo'lmasin)
  2. Savollar → DOCX formatida fayl yaratiladi
  3. Guruhga DOCX yuboriladi
  4. Faylga reply: test kartasi + Ulashish tugmasi
"""
import io
import logging
import os
import re
import tempfile

log = logging.getLogger(__name__)

LETTERS = list("ABCDEFGH")


# ═══════════════════════════════════════════════════════════════
# DOCX YARATISH
# ═══════════════════════════════════════════════════════════════

def _make_docx(questions: list, title: str = "", tid: str = "",
               category: str = "", creator_name: str = "") -> bytes:
    """
    Savollardan DOCX fayl yaratadi.
    Format:
        1. Savol matni
        A) variant
        *B) to'g'ri javob   ← yulduzcha
        Izoh: ...
    """
    try:
        from docx import Document
        from docx.shared import Pt, RGBColor, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        # python-docx yo'q — TXT bilan ketamiz
        return _make_txt(questions, title, tid, creator_name).encode("utf-8")

    doc = Document()

    # Sarlavha
    h = doc.add_heading(title or "Test", level=1)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # ── Meta ma'lumotlar (test savollariga halaqit qilmaydi) ──
    from datetime import datetime
    now_str = datetime.now().strftime("%d.%m.%Y %H:%M")

    sep = doc.add_paragraph()
    sep_run = sep.add_run("─" * 45)
    sep_run.font.size = Pt(9)
    sep_run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER

    meta_items = []
    if tid:          meta_items.append(f"Kod: {tid}")
    if category:     meta_items.append(f"Fan: {category}")
    if creator_name: meta_items.append(f"Muallif: {creator_name}")
    meta_items.append(f"Jami: {len(questions)} ta savol")
    meta_items.append(f"Sana: {now_str}")

    meta_p = doc.add_paragraph(" · ".join(meta_items))
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in meta_p.runs:
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
        run.italic = True

    sep2 = doc.add_paragraph()
    sep2_run = sep2.add_run("─" * 45)
    sep2_run.font.size = Pt(9)
    sep2_run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
    sep2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Savollar
    for i, q in enumerate(questions, 1):
        qt   = q.get("question") or q.get("text") or q.get("q") or ""
        opts = q.get("options", [])
        corr = q.get("correct", "")
        expl = q.get("explanation", "") or ""

        # To'g'ri javob indeksini aniqlash
        correct_idx = _resolve_correct_idx(q, opts)

        # Savol matni — qalin
        qp = doc.add_paragraph()
        qr = qp.add_run(f"{i}. {qt}")
        qr.bold = True
        qr.font.size = Pt(11)

        # Variantlar
        for j, opt in enumerate(opts):
            clean = re.sub(r'^[A-H]\s*[).]\s*', '', str(opt)).strip()
            lbl   = LETTERS[j] if j < len(LETTERS) else str(j)
            is_ok = (j == correct_idx)

            op = doc.add_paragraph()
            prefix = "*" if is_ok else " "
            or_ = op.add_run(f"  {prefix}{lbl}) {clean}")
            or_.font.size = Pt(10.5)
            if is_ok:
                or_.bold = True
                or_.font.color.rgb = RGBColor(0x00, 0x80, 0x00)  # yashil

        # Izoh
        if expl:
            ep = doc.add_paragraph()
            er = ep.add_run(f"  💡 Izoh: {expl}")
            er.italic = True
            er.font.size = Pt(9.5)
            er.font.color.rgb = RGBColor(0x44, 0x44, 0x88)

        doc.add_paragraph()  # bo'sh qator

    # Watermark
    try:
        _add_watermark(doc, creator_name=creator_name, bot_username="Quizmarkerbot")
    except Exception as _we:
        log.warning(f"Watermark xato: {_we}")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _make_txt(questions: list, title: str = "", tid: str = "",
              creator_name: str = "") -> str:
    """DOCX bo'lmasa TXT fallback"""
    lines = []
    if title:
        lines.append(f"📝 {title}")
    if tid:
        lines.append(f"🆔 {tid}")
    if creator_name:
        lines.append(f"👤 {creator_name}")
    lines.append(f"📋 {len(questions)} ta savol")
    lines.append("=" * 32)
    lines.append("")

    for i, q in enumerate(questions, 1):
        qt   = q.get("question") or q.get("text") or q.get("q") or ""
        opts = q.get("options", [])
        expl = q.get("explanation", "") or ""
        correct_idx = _resolve_correct_idx(q, opts)

        lines.append(f"{i}. {qt}")
        for j, opt in enumerate(opts):
            clean = re.sub(r'^[A-H]\s*[).]\s*', '', str(opt)).strip()
            lbl   = LETTERS[j] if j < len(LETTERS) else str(j)
            prefix = "*" if j == correct_idx else ""
            lines.append(f"{prefix}{lbl}) {clean}")
        if expl:
            lines.append(f"Izoh: {expl}")
        lines.append("")

    return "\n".join(lines)


def _make_txt_plain(questions: list) -> str:
    """
    Sof TXT format — meta sarlavhasiz, faqat savol+variant+izoh.
    Format:
        1. Savol
        A) variant
        B) variant
        *C) variant   ← to'g'ri javob
        D) variant
        Izoh : tushuntirish
    """
    lines = []
    for i, q in enumerate(questions, 1):
        qt   = q.get("question") or q.get("text") or q.get("q") or ""
        opts = q.get("options", [])
        expl = q.get("explanation", "") or ""
        correct_idx = _resolve_correct_idx(q, opts)

        lines.append(f"{i}. {qt}")
        for j, opt in enumerate(opts):
            clean = re.sub(r'^[A-H]\s*[).]\s*', '', str(opt)).strip()
            lbl   = LETTERS[j] if j < len(LETTERS) else str(j)
            prefix = "*" if j == correct_idx else ""
            lines.append(f"{prefix}{lbl}) {clean}")
        if expl:
            lines.append(f"Izoh : {expl}")
        lines.append("")

    return "\n".join(lines)



def _add_watermark(doc, creator_name: str = "", bot_username: str = "Quizmarkerbot"):
    """Har sahifaga diagonal watermark: @Bot nomi + Muallif ismi"""
    from lxml import etree

    wm1 = ("@" + bot_username) if bot_username else "@Quizmarkerbot"
    wm2 = ("Muallif: " + creator_name) if creator_name else ""

    row2_xml = ""
    if wm2:
        row2_xml = (
            "<w:p>"
            '<w:pPr><w:jc w:val="center"/></w:pPr>'
            "<w:r>"
            "<w:rPr>"
            '<w:color w:val="D8D8D8"/>'
            '<w:sz w:val="52"/><w:szCs w:val="52"/>'
            "</w:rPr>"
            "<w:t>" + wm2 + "</w:t>"
            "</w:r>"
            "</w:p>"
        )

    xml_str = (
        '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
        "<w:rPr><w:noProof/></w:rPr>"
        "<w:pict>"
        '<v:shape xmlns:v="urn:schemas-microsoft-com:vml"'
        ' xmlns:o="urn:schemas-microsoft-com:office:office"'
        ' xmlns:w10="urn:schemas-microsoft-com:office:word"'
        ' id="watermark1" type="#_x0000_t136"'
        ' style="position:absolute;margin-left:0;margin-top:0;'
        "width:520pt;height:220pt;z-index:-251654144;"
        "mso-position-horizontal:center;"
        "mso-position-horizontal-relative:margin;"
        "mso-position-vertical:center;"
        'mso-position-vertical-relative:margin"'
        ' fillcolor="#C8C8C8" stroked="f"'
        ' o:allowoverlap="t">'
        "<v:textbox>"
        "<w:txbxContent>"
        "<w:p>"
        '<w:pPr><w:jc w:val="center"/></w:pPr>'
        "<w:r>"
        "<w:rPr>"
        "<w:b/>"
        '<w:color w:val="C8C8C8"/>'
        '<w:sz w:val="96"/><w:szCs w:val="96"/>'
        "</w:rPr>"
        "<w:t>" + wm1 + "</w:t>"
        "</w:r>"
        "</w:p>"
        + row2_xml +
        "</w:txbxContent>"
        "</v:textbox>"
        '<w10:wrap xmlns:w10="urn:schemas-microsoft-com:office:word"'
        ' w10:anchorx="margin" w10:anchory="margin"/>'
        "</v:shape>"
        "</w:pict>"
        "</w:r>"
    )

    for section in doc.sections:
        hdr = section.header
        if not hdr.paragraphs:
            hdr.add_paragraph()
        para = hdr.paragraphs[0]
        para.clear()
        try:
            wm_el = etree.fromstring(xml_str)
            para._p.append(wm_el)
        except Exception:
            from docx.shared import Pt, RGBColor
            r = para.add_run(wm1 + ("  |  " + wm2 if wm2 else ""))
            r.font.size = Pt(7)
            r.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
            r.italic = True


def _resolve_correct_idx(q: dict, opts: list) -> int:
    """To'g'ri javob indeksini aniqlash"""
    corr = q.get("correct", "")
    if isinstance(corr, int):
        return corr

    # "correct_index" mavjudmi?
    ci = q.get("correct_index")
    if isinstance(ci, int):
        return ci

    if isinstance(corr, str):
        # "B) ..." formatida
        m = re.match(r'^([A-H])\s*[).]', corr.strip(), re.IGNORECASE)
        if m:
            return ord(m.group(1).upper()) - 65

        # Variant matni bilan solishtiramiz
        corr_clean = re.sub(r'^[A-H]\s*[).]\s*', '', corr).strip()
        for j, opt in enumerate(opts):
            opt_clean = re.sub(r'^[A-H]\s*[).]\s*', '', str(opt)).strip()
            if opt_clean == corr_clean or corr_clean in opt_clean or opt_clean in corr_clean:
                return j
    return 0


# ═══════════════════════════════════════════════════════════════
# ASOSIY FUNKSIYA
# ═══════════════════════════════════════════════════════════════

async def publish_to_baza(
    bot,
    tid: str,
    title: str,
    questions: list,
    creator_id: int,
    creator_name: str = "",
    bot_username: str = "",
    category: str = "",
    difficulty: str = "medium",
    passing_score: int = 60,
):
    """
    Baza Guruhiga e'lon qilish:
      1. DOCX fayl yuborish
      2. Faylga reply — test kartasi + Ulashish
    """
    try:
        from config import BAZA_GROUP_ID
        gid = int(BAZA_GROUP_ID or 0)
    except Exception:
        gid = 0

    if not gid:
        log.info("BAZA_GROUP_ID yo'q — chiqib ketdi")
        return

    try:
        from aiogram.types import BufferedInputFile
        from aiogram.utils.keyboard import InlineKeyboardBuilder
        from aiogram.types import InlineKeyboardButton

        diff_map = {
            "easy":   "🟢 Oson",
            "medium": "🟡 O'rtacha",
            "hard":   "🔴 Qiyin",
            "expert": "⚡ Ekspert",
        }
        diff_txt = diff_map.get(difficulty, "🟡 O'rtacha")
        qc       = len(questions)

        # ── 1. DOCX fayl tayyorlash ──
        # Fayl nomi — test nomidan, xavfsiz belgilar
        safe_name = re.sub(r'[\\/:*?"<>|]', '', title or tid).strip() or tid
        safe_name = safe_name[:60]  # Maksimal uzunlik

        try:
            docx_bytes = _make_docx(
                questions,
                title=title,
                tid=tid,
                category=category,
                creator_name=creator_name,
            )
            filename = f"{safe_name}.docx"
        except Exception as de:
            log.warning(f"DOCX xato, TXT ga o'tish: {de}")
            docx_bytes = _make_txt(questions, title, tid, creator_name).encode("utf-8")
            filename   = f"{safe_name}.txt"

        doc_file = BufferedInputFile(docx_bytes, filename=filename)

        # ── TXT fayl ham tayyorlaymiz (sof format, meta sarlavhasiz) ──
        txt_str   = _make_txt_plain(questions)
        txt_bytes = txt_str.encode("utf-8")
        txt_file  = BufferedInputFile(txt_bytes, filename=f"{safe_name}.txt")

        # Caption
        caption = (
            f"📄 <b>{title}</b>\n"
            f"━━━━━━━━━━━━━━━━━━━\n"
            f"🆔 <code>{tid}</code>\n"
            f"📁 {category or 'Boshqa'}\n"
            f"📊 {diff_txt}\n"
            f"📋 <b>{qc} ta savol</b>\n"
            f"🎯 O'tish: {passing_score}%\n"
            f"👤 {creator_name or "Noma'lum"}"
        )

        # ── 2. DOCX guruhga, TXT esa Storage kanalga (database) ──
        file_msg = await bot.send_document(
            chat_id=gid,
            document=doc_file,
            caption=caption,
        )
        try:
            from config import STORAGE_CHANNEL_ID
            storage_gid = int(STORAGE_CHANNEL_ID or 0)
        except Exception:
            storage_gid = 0

        if storage_gid:
            try:
                await bot.send_document(
                    chat_id=storage_gid,
                    document=txt_file,
                    caption=f"📦 {title} | {tid}",
                )
            except Exception as te:
                log.warning(f"TXT storage kanalga yuborilmadi: {te}")
        else:
            log.info("STORAGE_CHANNEL_ID yo'q — TXT yuborilmadi")

        # ── 3. Faylga reply — test kartasi + tugmalar ──
        # MUHIM: "Web test" va "Quiz Poll" — bot deep-link orqali
        # (?start=webtest_TID / ?start=poll_TID). Guruh xabarida
        # web_app tugmasi ishlamaydi (Telegram cheklovi), shuning
        # uchun bosilganda avtomatik bot shaxsiy chatiga o'tkaziladi.
        bu = bot_username or ""
        if not bu:
            try:
                me = await bot.get_me()
                bu = me.username
            except Exception:
                bu = ""

        bld = InlineKeyboardBuilder()
        if bu:
            bld.row(
                InlineKeyboardButton(text="🌐 Web test",
                    url=f"https://t.me/{bu}?start=webtest_{tid}"),
                InlineKeyboardButton(text="📊 Quiz Test ",
                    url=f"https://t.me/{bu}?start=poll_{tid}"),
            )
        bld.row(
            InlineKeyboardButton(
                text="📨 Testni ulashish",
                switch_inline_query=f"test_{tid}",
            )
        )

        card = (
            f"📌 <b>Yangi test!</b>\n"
            f"━━━━━━━━━━━━━━━━━━━\n"
            f"📝 <b>{title}</b>\n"
            f"📋 {qc} savol | {diff_txt} | {category or 'Boshqa'}\n"
            f"🆔 <code>{tid}</code>\n\n"
            f"👆 Fayl yuqorida\n"
            f"👇 Boshlash: Web Yoki Quiz"
        )

        await bot.send_message(
            chat_id=gid,
            text=card,
            reply_to_message_id=file_msg.message_id,
            reply_markup=bld.as_markup(),
        )

        log.info(f"✅ Baza publish: {tid} → guruh {gid}")

    except Exception as e:
        log.error(f"❌ publish_to_baza xato: {e}", exc_info=True)
