"""
📄 PARSER — TXT/PDF/DOCX fayllardan savollar ajratish

QO'LLAB-QUVVATLANADIGAN FORMATLAR:

FORMAT A — Standart (raqamli):
  1. Savol matni?
  *A) To'g'ri javob
  B) Xato
  C) Xato

FORMAT B — ==== separator (ZIP fayllar):
  Savol matni
  ====
  #To'g'ri javob     ← # bilan to'g'ri
  ====
  Xato javob
  ++++               ← keyingi savol

FORMAT C — ? = formati (PDF):
  ? Savol matni
  =To'g'ri           ← bo'sh joy yo'q = to'g'ri
  = Xato

FORMAT D — Jadval (Ko'p ustunli DOCX):
  Savol | To'g'ri javob | Muqobil | Muqobil

FORMAT E — Ha/Yo'q, Bo'sh joy to'ldirish, Erkin javob

FORMAT J — Markersiz, ketma-ket (to'g'ri javob aniqlanmaydi, PDF/TXT):
  Savol matni?
  Variant 1          ← qaysi to'g'ri ekani BELGISIZ
  Variant 2
  Variant 3
  Variant 4
  (natija: _marked=False, correct="" — bot AI/Seryalik/Admin so'raydi)
"""
import re, logging, os, subprocess, tempfile
from pathlib import Path

log = logging.getLogger(__name__)


# ═══════════════════════════════════════════════════════════
#  ASOSIY KIRISH NUQTASI
# ═══════════════════════════════════════════════════════════

def check_images_in_file(path: str) -> dict:
    """
    Fayl ichida rasm bor-yo'qligini tekshiradi.
    Qaytaradi: {"has_images": bool, "count": int, "type": "docx/pdf/none"}

    Botda foydalanuvchiga "Bu faylda N ta rasm bor" deyish uchun.
    """
    import os
    ext = os.path.splitext(path)[1].lower()
    result = {"has_images": False, "count": 0, "type": "none"}

    try:
        if ext == ".docx":
            from docx import Document
            try:
                doc = Document(path)
                cnt = len(doc.inline_shapes)
                # ZIP orqali ham tekshiramiz (NULL bo'lsa)
                if cnt == 0:
                    import zipfile
                    with zipfile.ZipFile(path) as z:
                        cnt = len([n for n in z.namelist()
                                   if 'media/image' in n and not n.endswith('/')])
                result = {"has_images": cnt > 0, "count": cnt, "type": "docx"}
            except Exception:
                # NULL xatosi — ZIP orqali
                import zipfile
                with zipfile.ZipFile(path) as z:
                    cnt = len([n for n in z.namelist()
                               if 'media/image' in n and not n.endswith('/')])
                result = {"has_images": cnt > 0, "count": cnt, "type": "docx"}

        elif ext == ".pdf":
            try:
                import fitz
                doc = fitz.open(path)
                cnt = sum(len(page.get_images()) for page in doc)
                doc.close()
                result = {"has_images": cnt > 0, "count": cnt, "type": "pdf"}
            except ImportError:
                # pdfplumber fallback
                import pdfplumber
                with pdfplumber.open(path) as pdf:
                    cnt = sum(len(p.images) for p in pdf.pages)
                result = {"has_images": cnt > 0, "count": cnt, "type": "pdf"}

        elif ext == ".doc":
            # DOC → DOCX convert qilib tekshiramiz
            converted = _convert_doc(path)
            if converted != path:
                return check_images_in_file(converted)

    except Exception as e:
        log.warning(f"check_images_in_file: {e}")

    return result


def parse_file(path: str) -> list:
    ext = Path(path).suffix.lower()
    try:
        # .doc → .docx konvertatsiya
        if ext == ".doc":
            path = _convert_doc(path)
            if not path:
                return []
            ext = ".docx"

        if ext == ".docx":
            return _parse_docx(path)
        elif ext == ".pdf":
            return _parse_pdf(path)
        elif ext == ".txt":
            return _parse_txt(path)
        elif ext in (".xlsx", ".xls", ".xlsm"):
            return _parse_xlsx(path)
        else:
            return []
    except Exception as e:
        log.error(f"parse_file xato ({ext}): {e}", exc_info=True)
        return []


def _convert_doc(path: str) -> str:
    """
    .doc → .docx konvertatsiya.
    1. LibreOffice (mavjud bo'lsa)
    2. python-docx2txt (fallback)
    3. antiword (fallback)
    """
    outdir = tempfile.mkdtemp()

    # 1. LibreOffice — eng yaxshi natija
    try:
        r = subprocess.run(
            ["libreoffice", "--headless", "--convert-to", "docx",
             path, "--outdir", outdir],
            capture_output=True, timeout=60
        )
        new = os.path.join(outdir, Path(path).stem + ".docx")
        if os.path.exists(new) and os.path.getsize(new) > 100:
            log.info(f"DOC → DOCX (LibreOffice): {Path(path).name}")
            return new
    except Exception as e:
        log.warning(f"LibreOffice yo'q yoki xato: {e}")

    # 2. python-docx to'g'ridan o'qib ko'ramiz (ba'zi .doc lar ishlaydi)
    try:
        import docx as _docx
        _docx.Document(path)  # Agar ishlasa — .doc aslida .docx ekan
        log.info(f"DOC to'g'ridan DOCX sifatida o'qildi: {Path(path).name}")
        return path
    except Exception:
        pass

    # 3. Matn sifatida o'qish (oxirgi fallback)
    try:
        import docx2txt
        txt = docx2txt.process(path)
        if txt and len(txt) > 50:
            txt_path = os.path.join(outdir, Path(path).stem + ".txt")
            with open(txt_path, "w", encoding="utf-8") as f:
                f.write(txt)
            log.info(f"DOC → TXT (docx2txt): {Path(path).name}")
            return txt_path
    except Exception:
        pass

    log.error(f"DOC convert qilinmadi: {path}")
    return ""


# ═══════════════════════════════════════════════════════════
#  DOCX PARSER — barcha jadval va paragraf formatlar
# ═══════════════════════════════════════════════════════════

def _parse_docx(path: str) -> list:
    try:
        from docx import Document
        doc = Document(path)
    except Exception as e:
        if 'NULL' in str(e):
            log.warning(f"DOCX NULL rel, ZIP fallback: {e}")
            return _parse_docx_via_zip(path)
        log.error(f"DOCX ochilmadi: {e}")
        return []

    # FORMAT B ni BIRINCHI tekshiramiz — ==== + # + ++++ eng aniq marker
    single = [t for t in doc.tables if 1 <= len(t.columns) <= 2]
    if single:
        tlines = []
        for t in single:
            for row in t.rows:
                tlines.append(row.cells[0].text.strip())
            tlines.append("+++++")
        if _is_eq_format(tlines):
            q = _parse_eq_hash(tlines)
            if q:
                return q

    # FORMAT D1/D2/D3: Ko'p ustunli jadvallar
    def _is_real_table(t):
        if len(t.rows) < 5:
            return False
        hdr = " ".join(c.text.strip().lower() for c in t.rows[0].cells)
        return any(k in hdr for k in ["savol","вопрос","topshiriq","question","test"]) or len(t.rows) >= 15

    multicol = [t for t in doc.tables if len(t.columns) >= 4 and _is_real_table(t)]
    if multicol:
        # FORMAT D3: A|B|C|D ustunli jadval
        hdr_upper = [c.text.strip().upper() for c in multicol[0].rows[0].cells]
        if any(h in ("A","B","C","D") for h in hdr_upper):
            q = _parse_abcd_table(multicol)
            if q:
                return q

        # FORMAT D1: Sarlavhali jadval (Savol|To'g'ri|Muqobil)
        q = _parse_table_multicol(multicol)
        if q:
            return q

    # FORMAT D2: Juft ustunli jadval
    double_col = [t for t in doc.tables if len(t.columns) >= 8 and _is_real_table(t)]
    if double_col:
        q = _parse_table_double_col(double_col)
        if q:
            return q

    # FORMAT E: 5x1 jadval (savol + 4 variant, har jadval = 1 savol)
    # ==== belgisi yo'q bo'lsa ishlaydi
    single_1col = [t for t in doc.tables
                   if len(t.columns) == 1 and 3 <= len(t.rows) <= 8]
    if single_1col:
        q = _parse_5x1_tables(single_1col)
        if q:
            return q

    # Paragraflardan
    lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    # FORMAT B: paragraflar ==== + # + ++++
    if _is_eq_format(lines):
        q = _parse_eq_hash(lines)
        if q:
            return q

    # FORMAT H: Rasmli savollar
    if doc.inline_shapes:
        try:
            q = _parse_docx_with_images(doc)
            if q:
                return q
        except Exception as _e:
            log.warning(f"_parse_docx_with_images: {_e}")

    # FORMAT G: Har paragraf = 1 savol (ichida \n bilan variantlar)
    if any('\n' in l for l in lines):
        q = _parse_paragraph_per_question(lines)
        if q:
            return q

    # FORMAT I: Har QATOR = alohida paragraf (raqamsiz, prefikssiz),
    # savol so'roq/ikki nuqta/bo'sh joy bilan, variantlar ketma-ket,
    # to'g'ri javob "+" bilan belgilangan (label A)/B)/C)/D) bo'lishi shart emas).
    # Paragraf ichida "\n" bo'lsa ham (FORMAT G mos kelmagan holatlar uchun)
    # avval qatorlarga yoyamiz, keyin FORMAT I logikasi ishlaydi.
    flat_lines = []
    for l in lines:
        flat_lines.extend(s.strip() for s in l.split('\n') if s.strip())
    q = _parse_loose_paragraphs(flat_lines)
    if q:
        return q

    # FORMAT F: Standart raqamli
    full_text = "\n".join(lines)
    q = parse_text(full_text)
    if q:
        return q

    # FORMAT F2: Variantsiz
    return _parse_question_list_only(lines)


def _parse_table_double_col(tables) -> list:
    """
    FORMAT D2: Juft ustunli jadval
    Har ustun 2 marta takrorlangan:
    | № | № | Savol | Savol | To'g'ri | To'g'ri | Muqobil | Muqobil |
    Faqat juft indeksli ustunlarni olamiz: 0,2,4,6,8...
    """
    LBL = ["A", "B", "C", "D", "E", "F", "G", "H"]
    questions = []
    for table in tables:
        if not table.rows or len(table.columns) < 6:
            continue
        # Header dan unique ustunlarni topamiz (har 2-ustun)
        header = [c.text.strip().lower() for c in table.rows[0].cells]
        # Juft indeksli ustunlar (0,2,4,6...)
        unique_cols = list(range(0, len(header), 2))

        q_col = -1; corr_col = -1; alt_cols = []
        for idx in unique_cols:
            h = header[idx] if idx < len(header) else ""
            if any(k in h for k in ["savol","topshiriq","вопрос"]):
                q_col = idx
            elif any(k in h for k in ["to'g'ri","tog'ri","правильн","correct","to`g`ri"]):
                corr_col = idx
            elif any(k in h for k in ["muqobil","variant","альт","incorrect"]):
                alt_cols.append(idx)

        if q_col == -1 or corr_col == -1:
            # Avtomatik: 2-ustun savol, 4-ustun to'g'ri, 6+ muqobil
            if len(unique_cols) >= 3:
                q_col = unique_cols[1] if len(unique_cols) > 1 else 2
                corr_col = unique_cols[2] if len(unique_cols) > 2 else 4
                alt_cols = [unique_cols[i] for i in range(3, min(len(unique_cols), 8))]
            else:
                continue

        for row in table.rows[1:]:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) <= max(q_col, corr_col):
                continue
            q_text  = cells[q_col] if q_col < len(cells) else ""
            correct = cells[corr_col] if corr_col < len(cells) else ""
            alts    = [cells[i] for i in alt_cols if i < len(cells) and cells[i]]
            if not q_text or not correct:
                continue
            # Takrorlanganlarni olib tashlaymiz
            if q_text == (cells[q_col+1] if q_col+1 < len(cells) else ""):
                pass  # Takrorlangan — normaldir
            all_opts = list(dict.fromkeys([correct] + alts))  # unique
            opts = []
            for i, o in enumerate(all_opts):
                lbl = LBL[i] if i < len(LBL) else str(i+1)
                opts.append(o if re.match(r"^[A-Ha-h]\s*[).]", o) else f"{lbl}) {o}")
            questions.append({
                "type":"multiple_choice","question":q_text,
                "options":opts,"correct":opts[0],
                "explanation":"","accepted_answers":[],"points":1,
                "_marked": True,
            })
    return questions


def _parse_abcd_table(tables) -> list:
    """
    FORMAT D3: A|B|C|D ustunli jadval
    Header: [FanBobi, Savol, A, B, C, D]
    To'g'ri javob belgilanmagan → _marked=False
    """
    LBL = ["A","B","C","D","E","F","G","H"]
    questions = []

    for table in tables:
        if len(table.columns) < 4 or not table.rows:
            continue
        header = [c.text.strip().upper() for c in table.rows[0].cells]

        # A, B, C, D ustunlarini topamiz
        opt_cols = [i for i, h in enumerate(header)
                    if h in ("A","B","C","D","E","F","G","H")]
        q_col    = next((i for i, h in enumerate(header)
                         if any(k in h.lower() for k in
                                ["SAVOL","TEST","TOPSHIRIG","ВОПРОС","TOPSHIRIQ"])), None)

        if not opt_cols:
            continue
        if q_col is None:
            # Avtomatik: birinchi A dan oldingi ustun = savol
            q_col = opt_cols[0] - 1 if opt_cols[0] > 0 else None
        if q_col is None or q_col >= len(header):
            continue
        if len(opt_cols) < 2:
            continue

        for row in table.rows[1:]:
            cells = [re.sub(r"\s+", " ", c.text).strip() for c in row.cells]
            if len(cells) <= max(q_col, max(opt_cols)):
                continue
            q_text = cells[q_col] if q_col < len(cells) else ""
            if not q_text or len(q_text) < 5:
                continue
            variants = [cells[ci] for ci in opt_cols if ci < len(cells) and cells[ci]]
            if len(variants) < 2:
                continue
            opts = []
            for i, v in enumerate(variants):
                lbl = LBL[i] if i < len(LBL) else str(i+1)
                opts.append(v if re.match(r"^[A-Ha-h]\s*[).]", v) else f"{lbl}) {v}")
            questions.append({
                "type": "multiple_choice", "question": q_text,
                "options": opts, "correct": "",  # Belgi yo'q — A variant TO'G'RI DEB OLINMAYDI
                "explanation": "", "accepted_answers": [], "points": 1,
                "_marked": False,
            })
    return questions


def _parse_5x1_tables(tables) -> list:
    """
    FORMAT E: 5x1 jadval — har jadval = 1 savol
    Qator 0: Savol matni
    Qator 1-4: Variantlar (to'g'ri javob belgilanmagan → differential)
    """
    LBL = ["A","B","C","D","E","F","G","H"]
    questions = []
    for table in tables:
        if len(table.columns) != 1:
            continue
        rows = [r.cells[0].text.strip() for r in table.rows if r.cells[0].text.strip()]
        if len(rows) < 3:
            continue
        q_text   = rows[0]
        variants = rows[1:]

        # Differential analysis bilan to'g'ri javobni topamiz
        raw = [f"= {v}" for v in variants]  # = prefix qo'shamiz
        clean_v, correct_idx = _find_correct_by_diff(raw)

        opts = []
        for i, v in enumerate(variants[:8]):
            lbl = LBL[i] if i < len(LBL) else str(i+1)
            opts.append(v if re.match(r"^[A-Ha-h]\s*[).]", v) else f"{lbl}) {v}")

        questions.append({
            "type":"multiple_choice","question":q_text,
            "options":opts,
            "correct":opts[correct_idx] if correct_idx >= 0 else "",
            "explanation":"","accepted_answers":[],"points":1,
            # Differensial tahlil orqali aniq topilgan bo'lsa belgilangan,
            # aks holda belgilanmagan — A variant TO'G'RI DEB OLINMAYDI
            "_marked": correct_idx >= 0,
        })
    return questions


def _parse_paragraph_per_question(lines: list) -> list:
    """
    FORMAT G: Har paragraf = 1 to'liq savol (raqam yo'q)

    Har element ichida \n bilan ajratilgan:
      Savol matni
      A) Variant 1
      B) ✅ To'g'ri variant
      C) Variant 3

    Har paragraf MUSTAQIL parse qilinadi.
    """
    LBL = ["A","B","C","D","E","F","G","H"]
    questions = []

    for raw_para in lines:
        # \n bo'lmasa — bu format emas
        if '\n' not in raw_para:
            continue

        sub = [l.strip() for l in raw_para.split('\n') if l.strip()]
        if len(sub) < 3:
            continue

        # Birinchi variant qatorini topamiz
        first_opt = None
        for idx, l in enumerate(sub):
            if re.match(r'^[A-Ha-h1-9]\s*[).]', l):
                first_opt = idx
                break

        if first_opt is None or first_opt == 0:
            continue

        # Savol matni
        q_text = ' '.join(sub[:first_opt]).strip()
        if not q_text:
            continue

        # Variantlar — FAQAT bu paragraf ichidagi
        opts = []
        correct_idx = -1

        for opt_line in sub[first_opt:]:
            vm = re.match(r'^([A-Ha-h1-9])\s*[).](.*)', opt_line)
            if not vm:
                continue

            label = vm.group(1).upper()
            rest  = vm.group(2).strip()
            is_correct = False

            # Boshida marker: ✅ * + # ...
            m_start = re.match(
                r'^([✅✓✔★☑•►→√■●▶◆*#+@]|={1,}[*+#]?|-{1,2})\s*(.+)',
                rest
            )
            if m_start:
                is_correct = True
                rest = m_start.group(2).strip()
            else:
                # Oxirida marker
                m_end = re.match(r'^(.+?)\s*([✅✓✔★☑•►→√■●▶◆*#+@])\s*$', rest)
                if m_end:
                    is_correct = True
                    rest = m_end.group(1).strip()

            if not rest:
                continue

            clean = f"{label}) {rest}"
            if is_correct and correct_idx == -1:
                correct_idx = len(opts)
            opts.append(clean)

        if len(opts) < 2:
            continue

        has_mark = correct_idx >= 0

        questions.append({
            "type":             "multiple_choice",
            "question":         q_text,
            "options":          opts,
            "correct":          opts[correct_idx] if has_mark else "",
            "explanation":      "",
            "accepted_answers": [],
            "points":           1,
            "_marked":          has_mark,
        })

    return questions


# ═══════════════════════════════════════════════════════════
#  FORMAT I: Har QATOR alohida paragraf, label/raqamsiz savol,
#  prefikssiz variantlar, to'g'ri javob "+" bilan belgilangan
#  (Word'da extract-text orqali olingan ko'rinish):
#
#    Savol matni?
#    Variant 1
#    +To'g'ri variant
#    Variant 3
#    Variant 4
#
#  Variantlar A)/B)/C)/D) prefiksli bo'lsa ham ishlaydi (aralash holatlar).
# ═══════════════════════════════════════════════════════════

def _looks_like_loose_question(line: str) -> bool:
    """
    Qator savol matniga o'xshaydimi — label/raqamsiz, variant emasligini tekshiradi.
    Belgilar: "?" yoki ":" bilan tugashi, bo'sh joy ifodalari (___, ------),
    yoki "N. " bilan boshlanishi (raqamli savol).
    """
    clean = line.replace('*', '').strip()
    if not clean:
        return False
    if clean.startswith('+'):
        return False
    if re.match(r'^[+]?[A-Ha-h]\s*[).]', clean):
        return False
    if clean.startswith('- '):
        inner = clean[2:].strip()
        if inner.startswith('+') or re.match(r'^[A-Ha-h]\s*[).]', inner):
            return False
    if clean.endswith('?') or clean.endswith(':'):
        return True
    if '___' in clean or '------' in clean or '______' in clean:
        return True
    if re.match(r'^\d+[.)]\s', clean):
        return True
    return False


def _clean_loose_option(line: str) -> tuple:
    """
    Bitta variant qatorini tozalaydi.
    Qaytaradi: (toza_matn, is_correct)
    """
    o = line.strip()
    o = o.replace('**', '').strip()
    if o.startswith('- '):
        o = o[2:].strip()
    is_correct = False
    if o.startswith('+'):
        is_correct = True
        o = o[1:].strip()
    # Allaqachon mavjud bo'lgan A)/B)/C)/D) prefiksini olib tashlaymiz
    # (keyinroq label qayta qo'yiladi — bir xil formatga keltirish uchun)
    m = re.match(r'^[+]?([A-Ha-h])\s*[).]\s*(.*)$', o)
    if m:
        o = m.group(2).strip()
    return o, is_correct


def _parse_loose_paragraphs(lines: list) -> list:
    """
    FORMAT I: Har savol va har variant — ALOHIDA, mustaqil paragraf
    (orasida \n yo'q — DOCX'da har bo'lim alohida paragraf bo'lganda).

    Mantiq:
      1. Savolga o'xshagan qatorni topamiz (_looks_like_loose_question)
      2. Undan keyingi ketma-ket Q-bo'lmagan qatorlarni variant deb olamiz
      3. Agar bir nechta Q-qator ketma-ket kelsa — bittasiga birlashtiramiz
         (masalan "Complete the sentence:" + "She ___ her arm...")
      4. "+" bilan boshlangan variant — to'g'ri javob
    """
    LBL = ["A", "B", "C", "D", "E", "F", "G", "H"]

    if not lines:
        return []

    tags = ['Q' if _looks_like_loose_question(l) else 'O' for l in lines]

    # Savol bo'lmagan birinchi qatorlar bo'lsa — tashlab yuboramiz
    items = list(lines)

    # Ketma-ket Q qatorlarni bitta savolga birlashtiramiz
    merged_items = []
    merged_tags = []
    i = 0
    n = len(items)
    while i < n:
        if tags[i] == 'Q':
            text = items[i]
            j = i + 1
            while j < n and tags[j] == 'Q':
                text = text + ' ' + items[j]
                j += 1
            merged_items.append(text)
            merged_tags.append('Q')
            i = j
        else:
            merged_items.append(items[i])
            merged_tags.append('O')
            i += 1

    questions = []
    n2 = len(merged_items)
    i = 0
    while i < n2:
        if merged_tags[i] != 'Q':
            i += 1
            continue

        q_text = merged_items[i].replace('**', '').strip()
        q_text = q_text[2:].strip() if q_text.startswith('- ') else q_text
        q_text = re.sub(r'^\d+[.)]\s*', '', q_text)

        j = i + 1
        variants = []
        correct_idx = -1
        while j < n2 and merged_tags[j] == 'O':
            opt_text, is_correct = _clean_loose_option(merged_items[j])
            if opt_text:
                if is_correct and correct_idx == -1:
                    correct_idx = len(variants)
                variants.append(opt_text)
            j += 1

        if q_text and len(variants) >= 2:
            opts = []
            for k, v in enumerate(variants):
                lbl = LBL[k] if k < len(LBL) else str(k + 1)
                opts.append(f"{lbl}) {v}")
            has_mark = correct_idx >= 0
            questions.append({
                "type":             "multiple_choice",
                "question":         q_text,
                "options":          opts,
                "correct":          opts[correct_idx] if has_mark else "",
                "explanation":      "",
                "accepted_answers": [],
                "points":           1,
                "_marked":          has_mark,
            })

        i = j if j > i else i + 1

    return questions


# ═══════════════════════════════════════════════════════════
#  FORMAT J: Savol/variant markersiz, to'g'ri javob BELGISIZ
#  (Test banki andozasi — PDF/TXT, ko'pincha sahifa raqamlari
#  va qator-wrap muammolari bilan birga keladi):
#
#    Savol matni?
#    Variant 1
#    Variant 2
#    Variant 3
#    Variant 4
#
#  Belgi (*, +, # va h.k.) UMUMAN YO'Q — qaysi variant to'g'ri
#  ekanini matndan aniqlash imkonsiz. Shuning uchun "correct"
#  BO'SH qoldiriladi va _marked=False qaytariladi — bot bu
#  savollarni "belgilanmagan" deb ko'rsatib, foydalanuvchiga
#  Seryalik javob / AI bilan yechish / Adminga murojaat
#  menyusini taqdim etadi (birinchi variantni "to'g'ri" deb
#  taxmin qilish XATO — bu chalg'ituvchi bo'lardi).
# ═══════════════════════════════════════════════════════════

def _strip_page_numbers(lines: list) -> list:
    """Mustaqil sahifa raqami qatorlarini (faqat 1-3 xonali raqam) olib tashlaydi."""
    return [l for l in lines if not re.match(r'^\d{1,3}$', l.strip())]


def _looks_like_marker_free_question(line: str) -> bool:
    """
    FORMAT J uchun savol qatorini aniqlaydi.
    Belgilar: "?" yoki ":" yoki "..." yoki "-" bilan tugashi,
    yoki bo'sh joy ifodasi (___, ...) borligi.
    (_looks_like_loose_question dan farqi — raqamli boshlanishni
    talab qilmaydi, chunki bu formatda savollar deyarli hech qachon
    "1." kabi raqamlanmaydi.)
    """
    clean = line.strip()
    if not clean:
        return False
    s = clean.rstrip()
    if s.endswith('?') or s.endswith(':') or s.endswith('...'):
        return True
    if s.endswith('-') and len(s) > 3:
        return True
    if '___' in s or '......' in s:
        return True
    return False


def _parse_marker_free_sequential(lines: list) -> list:
    """
    FORMAT J: Har savol va variant alohida qatorda, hech qanday
    to'g'ri-javob belgisi yo'q. Variantlar soni odatda 3-5 ta.

    Qaysi variant to'g'ri ekanini aniqlab bo'lmaydi — shuning uchun
    "correct" BO'SH ("") qoldiriladi, "_marked": False qaytariladi.
    Bot bunday savollarni "belgilanmagan" deb belgilab, Seryalik
    javob / AI bilan yechish / Adminga murojaat orqali to'ldirishni
    so'raydi.

    PDF dan kelgan matnda savol matni ba'zan 2 qatorga bo'linib
    qolishi mumkin (sahifa kengligi sababli) — bunday hollarda
    ketma-ket "savolga o'xshash" qatorlar bittasiga birlashtiriladi.
    """
    if not lines:
        return []

    items = _strip_page_numbers(lines)
    if not items:
        return []

    tags = ['Q' if _looks_like_marker_free_question(l) else 'O' for l in items]

    # Ketma-ket Q qatorlarni bitta savolga birlashtiramiz (PDF wrap holati)
    merged_items, merged_tags = [], []
    i, n = 0, len(items)
    while i < n:
        if tags[i] == 'Q':
            text = items[i]
            j = i + 1
            while j < n and tags[j] == 'Q':
                text = text + ' ' + items[j]
                j += 1
            merged_items.append(text)
            merged_tags.append('Q')
            i = j
        else:
            merged_items.append(items[i])
            merged_tags.append('O')
            i += 1

    questions = []
    n2 = len(merged_items)
    i = 0
    while i < n2:
        if merged_tags[i] != 'Q':
            i += 1
            continue

        q_text = merged_items[i].strip()

        j = i + 1
        variants = []
        while j < n2 and merged_tags[j] == 'O':
            v = merged_items[j].strip()
            if v:
                variants.append(v)
            j += 1

        # Kamida 2 ta variant bo'lishi shart (aks holda bu savol emas)
        if q_text and len(variants) >= 2:
            LBL = ["A", "B", "C", "D", "E", "F", "G", "H"]
            opts = []
            for k, v in enumerate(variants):
                lbl = LBL[k] if k < len(LBL) else str(k + 1)
                opts.append(f"{lbl}) {v}")
            questions.append({
                "type":             "multiple_choice",
                "question":         q_text,
                "options":          opts,
                "correct":          "",   # Belgi yo'q — A variant TO'G'RI DEB OLINMAYDI
                "explanation":      "",
                "accepted_answers": [],
                "points":           1,
                "_marked":          False,
            })

        i = j if j > i else i + 1

    return questions


def _parse_question_list_only(lines: list) -> list:
    """
    FORMAT F: Faqat savollar ro'yxati (variantsiz)
    Har qator = alohida savol. Variantlar yo'q → AI kerak.
    """
    questions = []
    for line in lines:
        line = line.strip()
        if len(line) < 10:
            continue
        q_text = re.sub(r"^\d+\s*[.)]\s*", "", line).strip()
        if not q_text:
            continue
        questions.append({
            "type":"multiple_choice","question":q_text,
            "options":["A) —","B) —","C) —","D) —"],
            "correct":"",  # Variantlarning o'zi yo'q — A) — TO'G'RI DEB OLINMAYDI
            "explanation":"","accepted_answers":[],"points":1,
            "_marked": False,
        })
    return questions


def _parse_table_multicol(tables) -> list:
    """Ko'p ustunli jadval: Savol | To'g'ri | Muqobil | Muqobil"""
    questions = []
    LBL = ["A", "B", "C", "D", "E", "F", "G", "H"]
    for table in tables:
        if not table.rows or len(table.columns) < 4:
            continue
        header = [c.text.strip().lower() for c in table.rows[0].cells]
        q_col = -1; corr_col = -1; alt_cols = []
        for i, h in enumerate(header):
            if any(k in h for k in ["savol", "topshiriq", "вопрос", "test"]):
                q_col = i
            elif any(k in h for k in ["to'g'ri", "tog'ri", "правильн", "correct"]):
                corr_col = i
            elif any(k in h for k in ["muqobil", "variant", "альт"]):
                alt_cols.append(i)
        if q_col == -1 or corr_col == -1:
            if len(header) >= 4:
                q_col = 1; corr_col = 2
                alt_cols = list(range(3, min(len(header), 7)))
            else:
                continue
        for row in table.rows[1:]:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) <= max(q_col, corr_col):
                continue
            q_text = cells[q_col]
            correct = cells[corr_col]
            alts = [cells[i] for i in alt_cols if i < len(cells) and cells[i]]
            if not q_text or not correct:
                continue
            all_opts = [correct] + alts
            opts = []
            for i, o in enumerate(all_opts):
                lbl = LBL[i] if i < len(LBL) else str(i+1)
                opts.append(o if re.match(r"^[A-Ha-h]\s*[).]", o) else f"{lbl}) {o}")
            questions.append({
                "type": "multiple_choice", "question": q_text,
                "options": opts, "correct": opts[0],
                "explanation": "", "accepted_answers": [], "points": 1,
                "_marked": True,
            })
    return questions


# ═══════════════════════════════════════════════════════════
#  FORMAT B: ==== + # + ++++ parser  (asosiy logika)
# ═══════════════════════════════════════════════════════════

def _is_eq_format(lines: list) -> bool:
    has_eq   = any(re.match(r"^={3,}$", l) for l in lines)
    has_plus = any(re.match(r"^\+{3,}$", l) for l in lines)
    return has_eq and has_plus


def _clean_text(text: str) -> str:
    """Matnni tozalaydi: bold, pipe, ko'p bo'shliq, bosh raqam"""
    import re as _re
    text = _re.sub(r'\*\*', '', text)
    text = _re.sub(r'\|', '', text)
    text = _re.sub(r'[ \t]+', ' ', text)
    text = _re.sub(r'^\d+[\.)\]]\s*', '', text)
    return text.strip()


def _is_valid_block(parts: list) -> bool:
    if not parts or len(parts[0]) > 800:
        return False
    return any(len(a.strip()) < 200 for a in parts[1:])


def _clean_table_block(block: str) -> str:
    lines = block.split('\n')
    if not any('|' in l for l in lines):
        return block
    cells = []
    for line in lines:
        line = line.strip()
        if not line.startswith('|'):
            if line: cells.append(line)
            continue
        if re.match(r'^\|[\s\-\|]+\|$', line):
            continue
        parts = [p.strip() for p in line.split('|') if p.strip()]
        for p in parts:
            if not re.match(r'^=+$', p):
                cells.append(p)
    return '\n'.join(cells)


def _parse_eq_hash(lines: list) -> list:
    """
    Asosiy logika (hujjatdan):
      ++++ → savollar chegarasi
      ==== → savol/javob chegarasi
      #    → to'g'ri javob belgisi
    """
    LBL = ["A", "B", "C", "D", "E", "F", "G", "H"]
    questions = []

    content = "\n".join(lines)

    # 1-QADAM: ++++ bo'yicha bloklarga ajratish
    blocks = re.split(r'\+{4,}', content)
    blocks = [b.strip() for b in blocks if b.strip()]

    for block in blocks:
        block = _clean_table_block(block)

        # 2-QADAM: ==== bo'yicha parts ga ajratish
        parts = re.split(r'={3,}', block)
        parts = [p.strip() for p in parts if p.strip()]

        if len(parts) < 2:
            continue
        if not _is_valid_block(parts):
            continue

        question = _clean_text(parts[0])
        if not question:
            continue

        # 3-QADAM: # bilan to'g'ri javobni topish
        correct_idx = -1
        clean_answers = []

        for ans in parts[1:]:
            ans = ans.strip()
            if ans.startswith('#'):
                if correct_idx == -1:
                    correct_idx = len(clean_answers)
                clean_answers.append(_clean_text(ans[1:]))
            else:
                clean_answers.append(_clean_text(ans))

        clean_answers = [a for a in clean_answers if a]
        if not clean_answers:
            continue

        has_mark = correct_idx != -1
        if has_mark and correct_idx >= len(clean_answers):
            # Xavfsizlik: marker topildi, lekin index chegaradan tashqari
            # chiqib ketgan bo'lsa (kutilmagan holat) — belgilanmagan deb olamiz
            has_mark = False
            correct_idx = -1

        opts = []
        for i, ans in enumerate(clean_answers):
            lbl = LBL[i] if i < len(LBL) else str(i + 1)
            opts.append(ans if re.match(r"^[A-Ha-h]\s*[).]", ans) else f"{lbl}) {ans}")

        questions.append({
            "type":             "multiple_choice",
            "question":         question,
            "options":          opts,
            "correct":          opts[correct_idx] if has_mark else "",
            "explanation":      "",
            "accepted_answers": [],
            "points":           1,
            "_marked":          has_mark,
        })

    return questions


# ═══════════════════════════════════════════════════════════
#  PDF PARSER
# ═══════════════════════════════════════════════════════════

def _parse_pdf(path: str) -> list:
    try:
        import pdfplumber
        pages = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    pages.append(t)
        full_text = "\n".join(pages)
    except Exception as e:
        log.error(f"PDF ochilmadi: {e}")
        return []

    lines = [l.strip() for l in full_text.split("\n") if l.strip()]

    # ─── MATNNI O'Z FORMATIDA PARSE QILAMIZ ───
    parsed = None
    q_cnt   = sum(1 for l in lines if l.startswith("?"))
    var_cnt = sum(1 for l in lines if any(l.startswith(v) for v in _VAR_STARTS))

    # FORMAT C: ? savol + = variant (differential)
    if q_cnt > 0 and var_cnt > q_cnt:
        parsed = _parse_question_eq(full_text)
    # FORMAT C2: ? savol, markersiz variantlar
    if not parsed and q_cnt > 0:
        parsed = _parse_question_no_marker(lines)
    # FORMAT B: ==== + ++++
    if not parsed and _is_eq_format(lines):
        parsed = _parse_eq_hash(lines)
    # FORMAT A: standart
    if not parsed:
        parsed = parse_text(full_text)
        # Sifat tekshiruvi: agar matn uzun bo'lib, juda kam savol
        # topilgan bo'lsa — bu FORMAT A mos kelmaganidan dalolat
        # beradi (masalan, raqamsiz FORMAT J matni "1 ta chiqindi
        # savol"ga aylanib qolishi mumkin). Bunday holda rad etamiz,
        # FORMAT J o'rniga ishlasin.
        if parsed and len(lines) > 30 and len(parsed) < max(3, len(lines) // 30):
            parsed = None
    # FORMAT J: savol/variant markersiz, ? yoki : bilan tugaydi,
    # birinchi variant = to'g'ri javob (test banki konventsiyasi).
    # ENG OXIRGI fallback — faqat boshqa hech qaysi format mos kelmasa
    # ishlatiladi (aks holda raqamli A)/B)/*/+ belgili savollarni
    # noto'g'ri ravishda tutib olib, to'g'ri javob belgisini yo'qotib
    # yuborishi mumkin).
    if not parsed:
        parsed = _parse_marker_free_sequential(lines)

    # ─── UNIVERSAL RASM BIRIKTIRISH ───
    # Format qanday bo'lishidan qat'i nazar, rasmlarni savollarga biriktiramiz
    if parsed:
        try:
            parsed = _attach_pdf_images(path, parsed, full_text)
        except Exception as _ie:
            log.warning(f"Rasm biriktirish: {_ie}")
        return parsed

    # Agar hech narsa topilmasa — bo'sh
    return parse_text(full_text)



# ═══════════════════════════════════════════════════════════
#  CORRECT MARKERS — Differential Analysis
# ═══════════════════════════════════════════════════════════
#
#  Mohiyat: variantlarni o'zaro taqqosla.
#  Qaysi variant BOSHQALARDAN FARQ QILSA — o'sha TO'G'RI.
#  Hammasi bir xil bo'lsa — BELGILANMAGAN.
#
#  Misol:
#    = Xato 1      marker: "= "
#    =+To'g'ri     marker: "=+"   ← FARQLI → TO'G'RI
#    = Xato 2      marker: "= "
#    = Xato 3      marker: "= "
# ═══════════════════════════════════════════════════════════

# Barcha taniqli variant belgilari (variant boshlanishi)
_VAR_STARTS = (
    '=', '+', '*', '#', '•', '►', '→', '✓', '✔', '√',
    '■', '●', '▶', '◆', '★', '☑', '–', '—',
)


def _extract_prefix(line: str) -> tuple:
    """
    Qatordan prefiks (marker) va toza matnni ajratadi.
    
    Qaytaradi: (prefix: str, clean: str) | (None, None) variant emas
    
    Misol:
      "=+To'g'ri"   → ("=+", "To'g'ri")
      "= Xato"      → ("= ", "Xato")
      "=Xato"       → ("=", "Xato")
      "*To'g'ri"    → ("*", "To'g'ri")
      "Oddiy matn"  → (None, None)
    """
    s = line.strip()
    if not s:
        return None, None
    
    # Variant belgisi bilan boshlanadimi?
    starts_with_marker = any(s.startswith(v) for v in _VAR_STARTS)
    if not starts_with_marker:
        return None, None
    
    # Prefiksni olish: harf yoki raqam kelguncha
    prefix = ""
    for ch in s:
        if ch.isalpha() or ch.isdigit():
            break
        prefix += ch
    
    clean = s[len(prefix):].strip()
    if not clean:
        return None, None
    
    return prefix, clean


def _find_correct_by_diff(raw_variants: list) -> tuple:
    """
    Differential analysis bilan to'g'ri javobni topadi.
    
    Qaytaradi: (variants: list[str], correct_idx: int)
      correct_idx = -1 → belgilanmagan
    
    Algoritm:
      1. Har variantdan prefiks ajrat
      2. Prefikslar orasidan ENG KAM UCHRAGANINI top
      3. Faqat 1 marta uchrasa → o'sha TO'G'RI
      4. Hammasi teng → BELGILANMAGAN
    """
    from collections import Counter
    
    parsed = []  # [(prefix, clean_text), ...]
    for raw in raw_variants:
        prefix, clean = _extract_prefix(raw)
        if prefix is None:
            continue
        parsed.append((prefix, clean))
    
    if not parsed:
        return [], -1
    
    # Prefikslarni sanab chiqamiz
    prefixes = [p for p, _ in parsed]
    counts = Counter(prefixes)
    
    correct_idx = -1
    
    if len(counts) > 1:
        # Eng kam uchraydigan prefiks
        min_count = min(counts.values())
        rarest = [p for p, c in counts.items() if c == min_count]
        
        # Faqat bitta farqli prefiks bo'lsa — o'sha to'g'ri
        if len(rarest) == 1:
            rare_prefix = rarest[0]
            for i, (p, _) in enumerate(parsed):
                if p == rare_prefix:
                    correct_idx = i
                    break
    
    # Toza variantlarni qaytaramiz
    clean_variants = [clean for _, clean in parsed]
    return clean_variants, correct_idx


def _parse_question_eq(text: str) -> list:
    """
    FORMAT C — ? savol, variantlar (differential analysis bilan):

    Qo'llab-quvvatlanadigan barcha holatlar:

    Holat 1 — Farqli marker (to'g'ri aniqladi):
      ? Savol
      = Xato 1
      =+ To'g'ri    ← =+ farqli → TO'G'RI
      = Xato 2
      = Xato 3

    Holat 2 — Bitta marker boshqacha:
      ? Savol
      = Variant 1
      * To'g'ri     ← * farqli → TO'G'RI
      = Variant 2
      = Variant 3

    Holat 3 — Hammasi bir xil (belgilanmagan):
      ? Savol
      = Variant 1
      = Variant 2
      = Variant 3
      = Variant 4

    Holat 4 — Eski format (bo'sh joysiz = to'g'ri):
      ? Savol
      = Xato 1
      =To'g'ri      ← bo'sh joy yo'q, boshqalari bo'sh joylik → FARQLI
      = Xato 2
    """
    LBL = ["A", "B", "C", "D", "E", "F", "G", "H"]
    questions = []
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    i = 0

    while i < len(lines):
        # Savol qatorini topamiz
        if not lines[i].startswith("?"):
            i += 1
            continue

        # Savol matni
        q_parts = [lines[i][1:].strip()]
        i += 1
        while i < len(lines):
            if lines[i].startswith("?"):
                break
            prefix, _ = _extract_prefix(lines[i])
            if prefix is not None:
                break
            q_parts.append(lines[i])
            i += 1

        q_text = " ".join(q_parts).strip()
        if not q_text:
            continue

        # Xom variantlarni yig'amiz
        raw_variants = []
        while i < len(lines) and not lines[i].startswith("?"):
            prefix, clean = _extract_prefix(lines[i])
            if prefix is not None:
                raw_variants.append(lines[i])
            i += 1

        if not raw_variants:
            continue

        # Differential analysis
        clean_variants, correct_idx = _find_correct_by_diff(raw_variants)

        if not clean_variants:
            continue

        # A) B) C) D) label qo'shish
        opts = []
        for j, v in enumerate(clean_variants):
            lbl = LBL[j] if j < len(LBL) else str(j + 1)
            opts.append(v if re.match(r"^[A-Ha-h]\s*[).]", v) else f"{lbl}) {v}")

        questions.append({
            "type":             "multiple_choice",
            "question":         q_text,
            "options":          opts,
            "correct":          opts[correct_idx] if correct_idx >= 0 else "",
            "explanation":      "",
            "accepted_answers": [],
            "points":           1,
            "_marked":          correct_idx >= 0,
        })

    return questions


# ═══════════════════════════════════════════════════════════
#  TXT PARSER
# ═══════════════════════════════════════════════════════════

def _parse_txt(path: str) -> list:
    text = _read_txt(path)
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    # FORMAT C: ? savol + = variant (differential)
    if sum(1 for l in lines if l.startswith("?")) > 0:
        # Avval differential analiz
        q = _parse_question_eq(text)
        if q:
            return q
        # FORMAT C2: ? savol, markersiz variant
        q = _parse_question_no_marker(lines)
        if q:
            return q
    # FORMAT B: ==== + # + ++++
    if _is_eq_format(lines):
        q = _parse_eq_hash(lines)
        if q:
            return q
    # FORMAT A: standart
    q = parse_text(text)
    if q and len(lines) > 30 and len(q) < max(3, len(lines) // 30):
        q = None
    if q:
        return q
    # FORMAT J: savol/variant markersiz, ? yoki : bilan tugaydi,
    # birinchi variant = to'g'ri javob (test banki konventsiyasi).
    # ENG OXIRGI fallback.
    return _parse_marker_free_sequential(lines)

def _read_txt(path: str) -> str:
    for enc in ("utf-8", "utf-8-sig", "cp1251", "latin-1"):
        try:
            with open(path, encoding=enc) as f:
                return f.read()
        except UnicodeDecodeError:
            pass
    return ""


# ═══════════════════════════════════════════════════════════
#  FORMAT A — STANDART RAQAMLI (mavjud logika saqlanadi)
# ═══════════════════════════════════════════════════════════

def parse_text(text: str) -> list:
    text = text.replace("\r\n", "\n")
    blocks = re.split(r"\n(?=\d+[\.)] *\S)", "\n" + text.strip())
    result = []
    for b in blocks:
        q = _parse_block(b.strip())
        if q:
            result.append(q)
    return result


def _clean_dup_label(text: str) -> str:
    """
    Variant boshidagi takroriy label ni tozalaydi.
    "A) a)Matn"  → "Matn"
    "B) b) Matn" → "Matn"
    "a)Matn"     → "Matn"
    Faqat MOS keluvchi takror (A)a)) yoki yakka kichik (a)) tozalanadi.
    """
    if not text:
        return text
    t = text.strip()
    # 1. Katta + kichik MOS label: "A) a)" yoki "A)a)" yoki "A. a."
    m = re.match(r'^([A-Za-z])\s*[).]\s*([a-z])\s*[).]\s*', t)
    if m and m.group(1).lower() == m.group(2).lower():
        return t[m.end():].strip()
    # 2. Yakka kichik label boshida: "a)Matn" (lekin katta emas)
    m2 = re.match(r'^([a-z])\s*[).]\s*(?=\S)', t)
    if m2:
        return t[m2.end():].strip()
    return t


def _is_correct_marker(line: str) -> tuple:
    """
    To'g'ri javob belgisini aniqlaydi — 3 ta holat:

    1. BOSHIDA marker:   *A) variant   +B) variant   ===C) variant
    2. LABEL+MARKER:     A)* variant   A)+ variant   A)# variant
    3. OXIRIDA marker:   A) variant*   A) variant +   A) variant#
    """
    ls = line.strip()
    if not ls:
        return False, ls

    _M = r'[*+#•►→✓✔√■●▶◆★☑\-=@~^]'

    # 1. BOSHIDA: ===... yoki marker+harf
    if ls.startswith("==="):
        return True, ls[3:].strip()

    # ==N (2+ teng belgi): ==To'g'ri, ==== ...
    m_eq = re.match(r'^={2,}\s*(.+)$', ls)
    if m_eq:
        return True, m_eq.group(1).strip()

    # --N (2+ tire): --To'g'ri, --- ...
    m_dash = re.match(r'^-{2,}\s*(.+)$', ls)
    if m_dash:
        return True, m_dash.group(1).strip()

    # =* =+ =# =- =• kombinatsiya (= + boshqa marker)
    m_comb = re.match(rf'^=[{re.escape("*+#•►→✓✔√■●▶◆★☑-@~^")}]\s*(.+)$', ls)
    if m_comb:
        return True, m_comb.group(1).strip()

    # marker + harf/raqam (bitta marker, keyin harf)
    if re.match(rf'^{_M}\s*[A-Za-zA-Яа-яёЁ0-9]', ls):
        return True, ls[1:].strip()

    # harfsiz: faqat marker + matn (masalan: * To'g'ri, + Javob)
    if re.match(rf'^{_M}\s+\S', ls):
        return True, ls[1:].strip()

    # 2. LABEL+MARKER O'RTADA: A)* variant, A)+ variant, A)# variant
    m = re.match(
        r'^([A-Ha-hА-Яа-яёЁ1-9])\s*[).]\s*'
        r'([*+#•►→✓✔√■●▶◆★☑\-=])\s*'
        r'(.+)$', ls
    )
    if m:
        label, marker, text = m.group(1), m.group(2), m.group(3).strip()
        if marker in ('-', '=') and not re.match(r'^[A-Za-zA-Яа-яёЁ0-9(«"\']', text):
            return False, ls
        return True, f"{label}) {text}"

    # 3. OXIRIDA: A) variant* yoki A) variant +
    m2 = re.match(
        r'^([A-Ha-hА-Яа-яёЁ1-9])\s*[).]\s*'
        r'(.+?)\s*'
        r'([*+#•►→✓✔√■●▶◆★☑=])\s*$', ls
    )
    if m2:
        label, text, marker = m2.group(1), m2.group(2).strip(), m2.group(3)
        if len(text) < 1:
            return False, ls
        return True, f"{label}) {text}"

    return False, ls


def _parse_block(block: str) -> dict | None:
    lines = [l.rstrip() for l in block.split("\n") if l.strip()]
    if not lines:
        return None

    forced = None
    if lines[0].upper().startswith("TYPE:"):
        forced = lines[0].split(":", 1)[1].strip().lower()
        lines = lines[1:]
    if not lines:
        return None

    q_text = re.sub(r"^\d+[\.)] *", "", lines[0]).strip()
    if not q_text:
        return None

    opts = []; corr = None; expl = ""; javob = None; acc = []; photo_id = None

    pm = re.match(r"^\[rasm:\s*([^\]]+)\]\s*", q_text)
    if pm:
        photo_id = pm.group(1).strip()
        q_text = q_text[pm.end():].strip()

    for line in lines[1:]:
        ls = line.strip()
        if not ls:
            continue
        if ls.startswith("[rasm:") and ls.endswith("]"):
            photo_id = ls[6:-1].strip()
            continue
        if ls.lower().startswith("izoh:"):
            expl = ls.split(":", 1)[1].strip()
            continue
        if re.match(r"^(qabul|accepted)\s*:", ls, re.IGNORECASE):
            acc = [a.strip() for a in re.split(r"[,;]", ls.split(":", 1)[1]) if a.strip()]
            continue
        if ls.lower().startswith("javob:"):
            javob = ls.split(":", 1)[1].strip()
            continue
        is_correct, cleaned = _is_correct_marker(ls)
        if is_correct:
            cleaned = _clean_dup_label(cleaned)
            opts.append(cleaned)
            corr = cleaned
            continue
        if re.match(r"^[A-Za-zA-Яа-яёЁ0-9]\s*[\).]\s*", ls):
            opts.append(_clean_dup_label(ls))
            continue

    if forced:
        qtype = forced
    elif javob is not None:
        jl = javob.lower().strip()
        if jl in ("ha", "yoq", "yo'q", "true", "false", "yes", "no"):
            qtype = "true_false"
        else:
            qtype = "fill_blank"
    elif opts:
        qtype = "multiple_choice"
    else:
        qtype = "text_input"

    if qtype == "true_false":
        corr = "Ha" if (javob or "").lower().strip() in ("ha", "true", "yes") else "Yo'q"
    elif qtype in ("text_input", "fill_blank"):
        corr = javob or corr or ""
    # multiple_choice uchun: agar hech qaysi variantda marker topilmagan
    # bo'lsa, corr ni BO'SH qoldiramiz (pastda has_marked=False bo'ladi) —
    # birinchi variantni "to'g'ri" deb OLINMAYDI.

    # Variantlardan markerlarni tozalaymiz
    # _is_correct_marker allaqachon toza matn qaytaradi
    # Lekin opts da hali marker qolgan bo'lishi mumkin
    def _strip_marker(text: str) -> str:
        ok, clean = _is_correct_marker(text)
        if ok:
            return clean
        # Oddiy variant — label ni saqlaymiz
        return re.sub(r"^===\s*", "", text).strip()

    clean_opts = [_strip_marker(o) for o in opts]
    if corr:
        corr = _strip_marker(corr)

    # multiple_choice/multi_select uchun: variant orasida * + # va h.k.
    # belgisi borligini tekshiramiz. true_false/fill_blank/text_input
    # uchun esa "Javob:" yorlig'i orqali javob berilgan-bermaganini
    # tekshiramiz — bu turlar uchun * # kabi belgi tabiatan ishlatilmaydi.
    if qtype in ("multiple_choice", "multi_select"):
        has_marked = any(_is_correct_marker(l)[0] for l in lines[1:] if l.strip())
    else:
        has_marked = bool(corr)

    result = {
        "type":             qtype,
        "question":         q_text,
        "options":          clean_opts if qtype in ("multiple_choice", "multi_select") else [],
        "correct":          corr or "",
        "explanation":      expl,
        "accepted_answers": acc,
        "points":           1,
        "_marked":          has_marked,
    }
    if photo_id:
        result["photo"] = photo_id
    return result


# ═══════════════════════════════════════════════════════════
# FORMAT C2: ? savol, markersiz variantlar (PDF)
# ═══════════════════════════════════════════════════════════

def _parse_question_no_marker(lines: list) -> list:
    """? savol, keyingi qatorlar variant (= markersiz)"""
    LBL = ["A","B","C","D","E","F","G","H"]
    questions = []
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line.startswith("?"):
            i += 1
            continue
        q_text = line[1:].strip()
        i += 1
        if not q_text:
            continue
        variants = []
        while i < len(lines):
            vl = lines[i].strip()
            if vl.startswith("?"):
                break
            if any(vl.startswith(v) for v in _VAR_STARTS):
                break
            if not vl:
                i += 1
                continue
            variants.append(vl)
            i += 1
        if len(variants) < 2:
            continue
        opts = []
        for j, v in enumerate(variants):
            lbl = LBL[j] if j < len(LBL) else str(j+1)
            opts.append(f"{lbl}) {v}" if not re.match(r"^[A-Ha-h]\s*[).]", v) else v)
        questions.append({
            "type": "multiple_choice", "question": q_text,
            "options": opts, "correct": "",  # Belgi yo'q — A variant TO'G'RI DEB OLINMAYDI
            "explanation": "", "accepted_answers": [], "points": 1,
            "_marked": False,
        })
    return questions


# ═══════════════════════════════════════════════════════════
# FORMAT X: XLSX parser
# ═══════════════════════════════════════════════════════════

def _parse_xlsx(path: str) -> list:
    """XLSX/XLS fayldan savollar"""
    try:
        import pandas as pd
    except ImportError:
        log.error("pandas o'rnatilmagan")
        return []
    try:
        xl = pd.ExcelFile(path)
    except Exception as e:
        log.error(f"XLSX ochilmadi: {e}")
        return []
    all_q = []
    for sheet in xl.sheet_names:
        try:
            df = pd.read_excel(xl, sheet_name=sheet, header=None)
            all_q.extend(_parse_xlsx_sheet(df))
        except Exception as e:
            log.warning(f"Sheet '{sheet}': {e}")
    return all_q


def _parse_xlsx_sheet(df) -> list:
    """Bitta XLSX sheet dan savollar"""
    import pandas as pd
    LBL = ["A","B","C","D","E","F","G","H"]
    Q_KW   = ["savol","вопрос","question","topshiriq"]
    C_KW   = ["to'g'ri","tog'ri","to`g`ri","правильн","correct","answer","togri"]
    BAD_KW = ["noto'g'ri","xato","incorrect","wrong","notogri"]

    header_row = None
    for ri in range(min(5, len(df))):
        row_vals = [str(v).strip().lower() for v in df.iloc[ri]
                    if str(v).strip() and str(v) != "nan"]
        if any(any(k in v for k in Q_KW) for v in row_vals):
            header_row = ri
            break

    header = [str(v).strip().lower() if str(v) != "nan" else ""
              for v in df.iloc[header_row]] if header_row is not None else []

    q_col = -1; corr_col = -1; opt_cols = []
    if header:
        for i, h in enumerate(header):
            if any(k in h for k in Q_KW): q_col = i
            elif any(k in h for k in C_KW) and not any(k in h for k in BAD_KW):
                corr_col = i
            elif any(k in h for k in BAD_KW) or h in ('a','b','c','d','e','f'):
                opt_cols.append(i)
    if q_col == -1: q_col = 0
    if corr_col == -1 and not opt_cols:
        corr_col = 1
        opt_cols = list(range(2, min(df.shape[1], 7)))

    data_start = (header_row + 1) if header_row is not None else 0
    questions  = []
    for ri in range(data_start, len(df)):
        row  = df.iloc[ri]
        vals = [str(v).strip() if (str(v) != "nan" and v == v) else "" for v in row]
        q_text = re.sub(r"^\d+[\.\)\t]\s*", "", vals[q_col] if q_col < len(vals) else "").strip()
        if not q_text or len(q_text) < 3: continue
        marked = False; correct = ""; variants = []
        if corr_col >= 0:
            correct = vals[corr_col] if corr_col < len(vals) else ""
            if not correct: continue
            is_c, clean = _is_correct_marker(correct)
            if is_c: correct = clean
            wrong = [vals[i] for i in opt_cols if i < len(vals) and vals[i]]
            variants = [correct] + wrong
            marked = True
        else:
            opt_pool = [vals[i] for i in opt_cols if i < len(vals) and vals[i]]
            if not opt_pool:
                opt_pool = [vals[i] for i in range(q_col+1, min(len(vals), q_col+6)) if vals[i]]
            cidx = -1; clean_pool = []
            for vi, v in enumerate(opt_pool):
                ic, cl = _is_correct_marker(v)
                if ic and cidx == -1: cidx = vi; clean_pool.append(cl)
                else: clean_pool.append(v)
            opt_pool = clean_pool
            if cidx >= 0:
                correct = opt_pool[cidx]
                variants = [opt_pool[cidx]] + [v for i,v in enumerate(opt_pool) if i != cidx]
                marked = True
            else:
                if not opt_pool: continue
                correct = ""; variants = opt_pool  # Belgi yo'q — birinchi variant TO'G'RI DEB OLINMAYDI
        if len(variants) < 2: continue
        opts = []
        for i, v in enumerate(variants):
            lbl = LBL[i] if i < len(LBL) else str(i+1)
            opts.append(v if re.match(r"^[A-Ha-h]\s*[).]", v) else f"{lbl}) {v}")
        questions.append({
            "type": "multiple_choice", "question": q_text,
            "options": opts, "correct": opts[0] if marked else "",
            "explanation": "", "accepted_answers": [], "points": 1,
            "_marked": marked,
        })
    return questions


# ═══════════════════════════════════════════════════════════
# FORMAT H: Rasmli DOCX parser
# ═══════════════════════════════════════════════════════════

def _parse_docx_with_images(doc) -> list:
    """# savol + inline rasm + - variantlar"""
    import os
    from docx.oxml.ns import qn as _qn
    LBL = ["A","B","C","D","E","F","G","H"]
    questions = []
    img_map = {}
    for i, p in enumerate(doc.paragraphs):
        blips = p._element.findall('.//' + _qn('a:blip'))
        if blips:
            for blip in blips:
                rId = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                if rId:
                    try:
                        part  = doc.part.related_parts[rId]
                        pname = str(part.partname)
                        if 'NULL' in pname.upper(): continue
                        img_map[i] = (part.blob, os.path.splitext(pname)[1].lower() or '.png')
                    except Exception:
                        pass
    if not img_map:
        return []
    paras = doc.paragraphs
    n = len(paras)
    i = 0
    while i < n:
        p    = paras[i]
        text = p.text.strip()
        if not (text.startswith('#') or (text and '#' in text[:5] and text[0] in ('#','\xa0'))):
            i += 1; continue
        q_text = text.lstrip('#\xa0').strip()
        i += 1
        img_bytes = img_ext = None
        if i < n and i in img_map:
            img_bytes, img_ext = img_map[i]; i += 1
        variants = []; correct_idx = -1
        while i < n:
            vp = paras[i]; vt = vp.text.strip()
            if vt.startswith('#') or (vt and '#' in vt[:5] and vt[0] in ('#','\xa0')): break
            if i in img_map and not vt: break
            if not vt: i += 1; continue
            if vt.startswith(('-','*','+','•')):
                is_c, clean = _is_correct_marker(vt)
                if vt.startswith('-') and not is_c: clean = vt[1:].strip()
                if is_c and correct_idx == -1: correct_idx = len(variants)
                variants.append(clean or vt[1:].strip())
            i += 1
        if not q_text or len(variants) < 2: continue
        opts = []
        for j, v in enumerate(variants):
            lbl = LBL[j] if j < len(LBL) else str(j+1)
            opts.append(v if re.match(r"^[A-Ha-h]\s*[).]", v) else f"{lbl}) {v}")
        has_mark = correct_idx >= 0
        q = {
            "type": "multiple_choice", "question": q_text,
            "options": opts, "correct": opts[correct_idx] if has_mark else "",
            "explanation": "", "accepted_answers": [], "points": 1,
            "_marked": has_mark, "_has_image": img_bytes is not None,
        }
        if img_bytes:
            q["_img_bytes"] = img_bytes; q["_img_ext"] = img_ext
        questions.append(q)
    return questions


def _parse_docx_via_zip(path: str) -> list:
    """NULL relationship xatosi bo'lganda ZIP orqali o'qiydi"""
    try:
        import zipfile
        from lxml import etree
    except ImportError:
        return []
    LBL = ["A","B","C","D","E","F","G","H"]
    questions = []
    try:
        with zipfile.ZipFile(path) as z:
            img_cache = {}
            try:
                rels_xml  = z.read('word/_rels/document.xml.rels')
                rels_root = etree.fromstring(rels_xml)
                for rel in rels_root:
                    rId    = rel.get('Id','')
                    target = rel.get('Target','')
                    rtype  = rel.get('Type','')
                    if 'image' in rtype and 'NULL' not in target.upper() and target:
                        fname = target.split('/')[-1]
                        try:
                            img_bytes = z.read(f'word/media/{fname}')
                            ext = os.path.splitext(fname)[1].lower() or '.png'
                            img_cache[rId] = (img_bytes, ext)
                        except Exception:
                            pass
            except Exception:
                pass
            doc_xml = z.read('word/document.xml')
            root    = etree.fromstring(doc_xml)
            W       = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
            R_NS    = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships'
            A_NS    = 'http://schemas.openxmlformats.org/drawingml/2006/main'
            paras   = root.findall(f'.//{{{W}}}p')
            n = len(paras); i = 0
            while i < n:
                p     = paras[i]
                texts = ''.join(t.text or '' for t in p.findall(f'.//{{{W}}}t'))
                text  = texts.strip()
                if not (text.startswith('#') or (text and '#' in text[:5])):
                    i += 1; continue
                q_text = text.lstrip('#\xa0').strip(); i += 1
                img_bytes = img_ext = None
                if i < n:
                    np = paras[i]
                    all_blips = np.findall('.//{*}blip')
                    if all_blips:
                        for blip in all_blips:
                            rId = blip.get(f'{{{R_NS}}}embed','')
                            if rId and rId in img_cache:
                                img_bytes, img_ext = img_cache[rId]; break
                        i += 1
                variants = []; correct_idx = -1
                while i < n:
                    vp   = paras[i]
                    vtxt = ''.join(t.text or '' for t in vp.findall(f'.//{{{W}}}t')).strip()
                    if vtxt.startswith('#') or (vtxt and '#' in vtxt[:5]): break
                    if not vtxt: i += 1; continue
                    if vtxt.startswith(('-','*','+','•')):
                        is_c, clean = _is_correct_marker(vtxt)
                        if vtxt.startswith('-') and not is_c: clean = vtxt[1:].strip()
                        if is_c and correct_idx == -1: correct_idx = len(variants)
                        variants.append(clean or vtxt[1:].strip())
                    i += 1
                if not q_text or len(variants) < 2: continue
                opts = []
                for j, v in enumerate(variants):
                    lbl = LBL[j] if j < len(LBL) else str(j+1)
                    opts.append(v if re.match(r"^[A-Ha-h]\s*[).]", v) else f"{lbl}) {v}")
                has_mark = correct_idx >= 0
                q = {
                    "type": "multiple_choice", "question": q_text,
                    "options": opts, "correct": opts[correct_idx] if has_mark else "",
                    "explanation": "", "accepted_answers": [], "points": 1,
                    "_marked": has_mark, "_has_image": img_bytes is not None,
                }
                if img_bytes:
                    q["_img_bytes"] = img_bytes; q["_img_ext"] = img_ext
                questions.append(q)
    except Exception as e:
        log.error(f"_parse_docx_via_zip: {e}")
    return questions



# ═══════════════════════════════════════════════════════════
# FORMAT P-IMG: Rasmli PDF parser (PyMuPDF)
# ═══════════════════════════════════════════════════════════

def _parse_pdf_with_images(path: str) -> list:
    """
    Rasmli PDF dan savollar + rasmlarni ajratadi (PyMuPDF).

    Har sahifada matn QATORLARI va rasmlar Y-koordinata bo'yicha
    aralashtiriladi. # bilan boshlangan qator = yangi savol.
    Savoldan keyingi rasm o'sha savolga biriktiriladi.
    """
    try:
        import fitz
    except ImportError:
        log.warning("PyMuPDF yo'q")
        return []

    LBL = ["A","B","C","D","E","F","G","H"]
    questions = []
    try:
        doc = fitz.open(path)
    except Exception as e:
        log.error(f"PDF (fitz) ochilmadi: {e}")
        return []

    # Elementlar: (page, y, kind, data)
    # kind: "line" (matn qatori) yoki "image"
    elements = []
    for page_no, page in enumerate(doc):
        # Matnni qator-qator olamiz (dict bilan aniqroq)
        try:
            pd = page.get_text("dict")
            for blk in pd.get("blocks", []):
                if blk.get("type") != 0:
                    continue
                for line in blk.get("lines", []):
                    spans = line.get("spans", [])
                    txt = "".join(sp.get("text","") for sp in spans).strip()
                    if txt:
                        y = line.get("bbox", [0,0,0,0])[1]
                        elements.append((page_no, y, "line", txt))
        except Exception:
            # Fallback: blocks
            for b in page.get_text("blocks"):
                if b[6] == 0 and b[4].strip():
                    for ln in b[4].strip().split("\n"):
                        if ln.strip():
                            elements.append((page_no, b[1], "line", ln.strip()))

        # Rasmlar
        for img_info in page.get_images(full=True):
            xref = img_info[0]
            try:
                bbox = page.get_image_bbox(img_info)
                y = bbox.y0
            except Exception:
                y = 999999
            try:
                base = doc.extract_image(xref)
                elements.append((page_no, y, "image",
                                 (base["image"], "." + base.get("ext","png"))))
            except Exception:
                pass

    doc.close()
    elements.sort(key=lambda e: (e[0], e[1]))

    # Savol belgisi: # yoki ?  | Variant belgisi: - * + • yoki =
    Q_MARK = ("#", "?")
    V_MARK = ("-", "*", "+", "•", "=")

    # Savol-rasm-variant ketma-ketligi
    i = 0
    n = len(elements)
    while i < n:
        _, _, kind, data = elements[i]
        if kind == "line" and data.startswith(Q_MARK):
            q_text = data.lstrip("#?\xa0 ").strip()
            i += 1
            img_bytes = img_ext = None
            variants = []
            correct_idx = -1
            while i < n:
                _, _, k2, d2 = elements[i]
                if k2 == "line" and d2.startswith(Q_MARK):
                    break
                if k2 == "image" and img_bytes is None:
                    img_bytes, img_ext = d2
                    i += 1
                    continue
                if k2 == "line":
                    line = d2.strip()
                    if line.startswith(V_MARK):
                        is_c, clean = _is_correct_marker(line)
                        # - yoki = bilan boshlanса va marker emas — oddiy variant
                        if line[0] in ("-", "=") and not is_c:
                            clean = line[1:].strip()
                        if is_c and correct_idx == -1:
                            correct_idx = len(variants)
                        variants.append(clean or line[1:].strip())
                    elif variants:
                        # Variant davomi (uzun variant)
                        variants[-1] += " " + line
                    else:
                        # Savol davomi
                        q_text += " " + line
                i += 1

            if q_text and len(variants) >= 2:
                opts = []
                for j, v in enumerate(variants):
                    lbl = LBL[j] if j < len(LBL) else str(j+1)
                    opts.append(v if re.match(r"^[A-Ha-h]\s*[).]", v) else f"{lbl}) {v}")
                has_mark = correct_idx >= 0
                q = {
                    "type": "multiple_choice", "question": q_text,
                    "options": opts, "correct": opts[correct_idx] if has_mark else "",
                    "explanation": "", "accepted_answers": [], "points": 1,
                    "_marked": has_mark, "_has_image": img_bytes is not None,
                }
                if img_bytes:
                    q["_img_bytes"] = img_bytes
                    q["_img_ext"] = img_ext
                questions.append(q)
        else:
            i += 1

    return questions

def _pdf_via_docx(path: str) -> list:
    """
    PDF ni LibreOffice bilan DOCX ga aylantiradi,
    keyin DOCX rasm parserini (_parse_docx_with_images) ishlatadi.

    LibreOffice rasm-matn joylashuvini PyMuPDF dan yaxshiroq saqlaydi.
    """
    import subprocess, tempfile, os
    from pathlib import Path

    out_dir = tempfile.mkdtemp(prefix="pdf2docx_")
    try:
        # LibreOffice bilan PDF → DOCX
        subprocess.run(
            ["libreoffice", "--headless", "--infilter=writer_pdf_import",
             "--convert-to", "docx", path, "--outdir", out_dir],
            check=True, timeout=120,
            stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
        )
        # Yaratilgan DOCX ni topamiz
        docx_files = list(Path(out_dir).glob("*.docx"))
        if not docx_files:
            log.warning("PDF→DOCX: DOCX yaratilmadi")
            return []

        docx_path = str(docx_files[0])
        log.info(f"PDF→DOCX muvaffaqiyatli: {Path(docx_path).name}")

        # DOCX rasm parserini ishlatamiz
        try:
            from docx import Document
            doc = Document(docx_path)
            q = _parse_docx_with_images(doc)
            if q:
                return q
            # Rasmli topilmasa — oddiy DOCX parser
            return _parse_docx(docx_path)
        except Exception as e:
            if 'NULL' in str(e):
                return _parse_docx_via_zip(docx_path)
            log.warning(f"PDF→DOCX parse xato: {e}")
            return []

    except subprocess.TimeoutExpired:
        log.warning("PDF→DOCX: timeout (120s)")
        return []
    except FileNotFoundError:
        log.warning("PDF→DOCX: LibreOffice topilmadi")
        return []
    except Exception as e:
        log.warning(f"PDF→DOCX xato: {e}")
        return []
    finally:
        # Vaqtinchalik papkani tozalaymiz (DOCX o'qilgandan keyin)
        # Lekin _img_bytes allaqachon xotirada, fayl kerak emas
        try:
            import shutil
            shutil.rmtree(out_dir, ignore_errors=True)
        except Exception:
            pass


# ═══════════════════════════════════════════════════════════
# UNIVERSAL RASM BIRIKTIRISH (barcha PDF formatlar uchun)
# ═══════════════════════════════════════════════════════════

def _attach_pdf_images(path: str, questions: list, full_text: str) -> list:
    """
    PDF dagi rasmlarni savollarga biriktiradi — FORMAT FARQI YO'Q.

    Mantiq:
      1. PDF dan barcha rasmlarni (bytes + Y-koordinata) ajratamiz
      2. Har savol matnining PDF dagi Y-pozitsiyasini topamiz
      3. Har rasmni — undan YUQORIDAGI eng yaqin savolga biriktiramiz
         (rasm odatda savol matnidan keyin keladi)

    Bu # , ? , ==== — barcha formatlarda bir xil ishlaydi.
    """
    try:
        import fitz
    except ImportError:
        # PyMuPDF yo'q — rasm biriktirib bo'lmaydi (lekin matn ishlaydi)
        return questions

    try:
        doc = fitz.open(path)
    except Exception as e:
        log.warning(f"_attach_pdf_images fitz: {e}")
        return questions

    # ─── 1. Rasmlarni ajratamiz (global Y bilan) ───
    images = []  # (global_y, img_bytes, img_ext)
    page_y_offset = 0.0
    for page in doc:
        page_h = page.rect.height
        for img_info in page.get_images(full=True):
            xref = img_info[0]
            try:
                bbox = page.get_image_bbox(img_info)
                gy = page_y_offset + bbox.y0
            except Exception:
                gy = page_y_offset
            try:
                base = doc.extract_image(xref)
                images.append((gy, base["image"], "." + base.get("ext", "png")))
            except Exception:
                pass
        page_y_offset += page_h

    if not images:
        doc.close()
        return questions

    # ─── 2. Savol matnlarining global Y pozitsiyasi ───
    # Har savolning birinchi 15 belgisini PDF dan qidiramiz
    q_positions = []  # (global_y, q_index)
    page_y_offset = 0.0
    page_texts = []
    for page in doc:
        page_texts.append((page_y_offset, page))
        page_y_offset += page.rect.height

    def _find_q_y(q_text):
        """Savol matnining global Y koordinatasini topadi"""
        # Savol matnidan qidiruv uchun bo'lak (markerlarsiz, qisqa)
        probe = re.sub(r'^[#?=*+\-•\s]+', '', q_text).strip()[:20]
        if len(probe) < 4:
            return None
        for y_off, page in page_texts:
            try:
                rects = page.search_for(probe, quads=False)
                if rects:
                    return y_off + rects[0].y0
            except Exception:
                continue
        return None

    for qi, q in enumerate(questions):
        gy = _find_q_y(q.get("question", ""))
        if gy is not None:
            q_positions.append((gy, qi))

    doc.close()

    if not q_positions:
        return questions

    q_positions.sort()

    # ─── 3. Har rasmni eng yaqin (yuqoridagi) savolga biriktiramiz ───
    for img_y, img_bytes, img_ext in images:
        # Rasmdan YUQORIDA va eng yaqin savolni topamiz
        best_qi = None
        best_dist = float("inf")
        for q_y, qi in q_positions:
            if q_y <= img_y:  # Savol rasmdan yuqorida
                dist = img_y - q_y
                if dist < best_dist:
                    best_dist = dist
                    best_qi = qi
        # Agar yuqorida savol topilmasa — eng yaqin pastdagisi
        if best_qi is None:
            for q_y, qi in q_positions:
                dist = abs(q_y - img_y)
                if dist < best_dist:
                    best_dist = dist
                    best_qi = qi

        # Biriktiramiz (agar bu savolda hali rasm bo'lmasa)
        if best_qi is not None and not questions[best_qi].get("_img_bytes"):
            questions[best_qi]["_img_bytes"] = img_bytes
            questions[best_qi]["_img_ext"]   = img_ext
            questions[best_qi]["_has_image"]  = True

    img_attached = sum(1 for q in questions if q.get("_img_bytes"))
    log.info(f"PDF rasm biriktirildi: {img_attached} ta savol")
    return questions
